import os
import asyncio
from typing import Optional
from pathlib import Path
import mimetypes
import logging

# 文本提取相关库
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    import speech_recognition as sr
except ImportError:
    sr = None

logger = logging.getLogger(__name__)

class TextExtractionService:
    """文本提取服务，支持多种文件格式"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._extract_from_pdf,
            'text/plain': self._extract_from_text,
            'text/csv': self._extract_from_csv,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_from_docx,
            'application/msword': self._extract_from_doc,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._extract_from_excel,
            'application/vnd.ms-excel': self._extract_from_excel,
            # 移除图片和音频处理，避免网络连接问题
            # 'image/jpeg': self._extract_from_image,
            # 'image/png': self._extract_from_image,
            # 'image/gif': self._extract_from_image,
            # 'image/bmp': self._extract_from_image,
            # 'image/tiff': self._extract_from_image,
            # 'audio/mpeg': self._extract_from_audio,
            # 'audio/wav': self._extract_from_audio,
            # 'audio/ogg': self._extract_from_audio,
        }
    
    async def extract_text(self, file_path: str, content_type: str) -> Optional[str]:
        """
        从文件中提取文本内容
        
        Args:
            file_path: 文件路径
            content_type: MIME类型
            
        Returns:
            提取的文本内容，如果失败返回None
        """
        try:
            # 检查是否为图片或音频文件
            if content_type.startswith('image/') or content_type.startswith('audio/'):
                logger.info(f"Skipping text extraction for media file: {content_type}")
                return None
            
            if content_type not in self.supported_types:
                logger.warning(f"Unsupported content type: {content_type}")
                return None
            
            extractor = self.supported_types[content_type]
            
            # 在线程池中运行提取操作，避免阻塞
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, extractor, file_path)
            
            return text
            
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """从PDF文件提取文本"""
        if not pdfplumber:
            logger.warning("pdfplumber not available for PDF extraction")
            return None
        
        try:
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            return '\n\n'.join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            # 尝试使用PyPDF2作为备选
            if PyPDF2:
                try:
                    text_content = []
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page in pdf_reader.pages:
                            text = page.extract_text()
                            if text:
                                text_content.append(text)
                    
                    return '\n\n'.join(text_content) if text_content else None
                except Exception as e2:
                    logger.error(f"PyPDF2 extraction also failed: {str(e2)}")
            
            return None
    
    def _extract_from_text(self, file_path: str) -> Optional[str]:
        """从纯文本文件提取内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ['gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
        except Exception as e:
            logger.error(f"Text file extraction failed: {str(e)}")
        
        return None
    
    def _extract_from_csv(self, file_path: str) -> Optional[str]:
        """从CSV文件提取内容"""
        if not pd:
            logger.warning("pandas not available for CSV extraction")
            return None
        
        try:
            df = pd.read_csv(file_path)
            # 将DataFrame转换为文本格式
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"CSV extraction failed: {str(e)}")
            return None
    
    def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """从DOCX文件提取文本"""
        if not Document:
            logger.warning("python-docx not available for DOCX extraction")
            return None
        
        try:
            # 检查文件是否存在和可读
            if not os.path.exists(file_path):
                logger.error(f"DOCX file not found: {file_path}")
                return None
            
            if not os.access(file_path, os.R_OK):
                logger.error(f"DOCX file not readable: {file_path}")
                return None
            
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            logger.info(f"Processing DOCX file: {file_path}, size: {file_size} bytes")
            
            # 尝试打开文档
            doc = Document(file_path)
            text_content = []
            
            # 提取段落内容
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    paragraph_count += 1
            
            logger.info(f"Extracted {paragraph_count} paragraphs from DOCX")
            
            # 提取表格内容
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            logger.info(f"Extracted {table_count} tables from DOCX")
            
            result = '\n\n'.join(text_content) if text_content else None
            if result:
                logger.info(f"Successfully extracted {len(result)} characters from DOCX")
            else:
                logger.warning("No text content found in DOCX file")
            
            return result
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"DOCX extraction failed for {file_path}: {error_msg}")
            
            # 检查是否是文件损坏的常见错误
            if "There is no item named" in error_msg or "archive" in error_msg.lower():
                logger.warning("DOCX file appears to have corrupted references, trying alternative extraction method")
                # 尝试备用提取方法
                return self._extract_from_docx_fallback(file_path)
            elif "Permission denied" in error_msg:
                logger.error("Permission denied when accessing DOCX file")
            elif "No such file" in error_msg:
                logger.error("DOCX file not found")
            
            return None
    
    def _extract_from_docx_fallback(self, file_path: str) -> Optional[str]:
        """备用DOCX文本提取方法，直接从ZIP文件中提取XML内容"""
        import zipfile
        import xml.etree.ElementTree as ET
        import re
        
        try:
            logger.info(f"Using fallback extraction method for DOCX: {file_path}")
            
            text_content = []
            
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # 提取主文档内容
                if 'word/document.xml' in zip_file.namelist():
                    document_xml = zip_file.read('word/document.xml').decode('utf-8')
                    text_content.extend(self._extract_text_from_xml(document_xml))
                
                # 提取页眉页脚内容（跳过有问题的文件）
                for file_name in zip_file.namelist():
                    if file_name.startswith('word/header') and file_name.endswith('.xml'):
                        try:
                            header_xml = zip_file.read(file_name).decode('utf-8')
                            text_content.extend(self._extract_text_from_xml(header_xml))
                        except Exception as e:
                            logger.warning(f"Skipping problematic header file {file_name}: {e}")
                    
                    elif file_name.startswith('word/footer') and file_name.endswith('.xml'):
                        try:
                            footer_xml = zip_file.read(file_name).decode('utf-8')
                            text_content.extend(self._extract_text_from_xml(footer_xml))
                        except Exception as e:
                            logger.warning(f"Skipping problematic footer file {file_name}: {e}")
            
            result = '\n\n'.join(text_content) if text_content else None
            if result:
                logger.info(f"Fallback extraction successful: {len(result)} characters extracted")
            else:
                logger.warning("Fallback extraction found no text content")
            
            return result
            
        except Exception as e:
            logger.error(f"Fallback DOCX extraction failed: {str(e)}")
            return None
    
    def _extract_text_from_xml(self, xml_content: str) -> list:
        """从XML内容中提取文本，正确处理段落结构"""
        import xml.etree.ElementTree as ET
        import re
        
        try:
            # 移除XML命名空间前缀以简化处理
            xml_content = re.sub(r'<w:', '<', xml_content)
            xml_content = re.sub(r'</w:', '</', xml_content)
            xml_content = re.sub(r'xmlns:w="[^"]*"', '', xml_content)
            xml_content = re.sub(r'w:', '', xml_content)
            
            root = ET.fromstring(xml_content)
            
            # 按段落提取文本
            paragraphs = []
            paragraph_elements = root.findall('.//p')
            
            logger.info(f"Found {len(paragraph_elements)} paragraphs in XML")
            
            for p_elem in paragraph_elements:
                # 提取段落中的所有文本元素
                text_elements = p_elem.findall('.//t')
                paragraph_text = []
                
                for t_elem in text_elements:
                    if t_elem.text:
                        paragraph_text.append(t_elem.text)
                
                # 合并段落文本
                if paragraph_text:
                    full_paragraph = ''.join(paragraph_text).strip()
                    if full_paragraph:
                        paragraphs.append(full_paragraph)
            
            logger.info(f"Extracted {len(paragraphs)} non-empty paragraphs")
            return paragraphs
            
        except Exception as e:
            logger.warning(f"XML text extraction failed: {e}")
            # 如果XML解析失败，尝试简单的正则表达式提取
            try:
                # 先尝试按段落提取
                paragraph_matches = re.findall(r'<p[^>]*>(.*?)</p>', xml_content, re.DOTALL)
                paragraphs = []
                for p_match in paragraph_matches:
                    text_matches = re.findall(r'<t[^>]*>([^<]+)</t>', p_match)
                    if text_matches:
                        paragraph_text = ''.join(text_matches).strip()
                        if paragraph_text:
                            paragraphs.append(paragraph_text)
                
                if paragraphs:
                    logger.info(f"Regex extraction found {len(paragraphs)} paragraphs")
                    return paragraphs
                else:
                    # 如果段落提取失败，直接提取所有文本
                    text_matches = re.findall(r'<t[^>]*>([^<]+)</t>', xml_content)
                    return [match.strip() for match in text_matches if match.strip()]
                    
            except Exception as regex_error:
                logger.error(f"Regex extraction also failed: {regex_error}")
                return []
    
    def _extract_from_doc(self, file_path: str) -> Optional[str]:
        """从DOC文件提取文本（需要额外工具）"""
        logger.warning("DOC format extraction not implemented")
        return None
    
    def _extract_from_excel(self, file_path: str) -> Optional[str]:
        """从Excel文件提取内容"""
        if not pd:
            logger.warning("pandas not available for Excel extraction")
            return None
        
        try:
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheet_text = f"工作表: {sheet_name}\n{df.to_string(index=False)}"
                text_content.append(sheet_text)
            
            return '\n\n'.join(text_content) if text_content else None
            
        except Exception as e:
            logger.error(f"Excel extraction failed: {str(e)}")
            return None
    
    def _extract_from_image(self, file_path: str) -> Optional[str]:
        """从图像文件提取文本（OCR）"""
        if not Image or not pytesseract:
            logger.warning("PIL or pytesseract not available for image OCR")
            return None
        
        try:
            image = Image.open(file_path)
            # 使用中英文OCR
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            return text.strip() if text.strip() else None
            
        except Exception as e:
            logger.error(f"Image OCR failed: {str(e)}")
            return None
    
    def _extract_from_audio(self, file_path: str) -> Optional[str]:
        """从音频文件提取文本（语音识别）"""
        if not sr:
            logger.warning("speech_recognition not available for audio transcription")
            return None
        
        try:
            r = sr.Recognizer()
            
            # 支持WAV格式
            if file_path.lower().endswith('.wav'):
                with sr.AudioFile(file_path) as source:
                    audio = r.record(source)
                    # 尝试使用Google语音识别
                    text = r.recognize_google(audio, language='zh-CN')
                    return text
            else:
                logger.warning(f"Audio format not supported for transcription: {file_path}")
                return None
                
        except sr.UnknownValueError:
            logger.warning("Speech recognition could not understand audio")
            return None
        except sr.RequestError as e:
            logger.error(f"Speech recognition service error: {str(e)}")
            return None
        except Exception as e:
            logger.error(f"Audio transcription failed: {str(e)}")
            return None
    
    def get_supported_types(self) -> list:
        """获取支持的文件类型列表"""
        return list(self.supported_types.keys())
    
    def is_supported(self, content_type: str) -> bool:
        """检查是否支持指定的文件类型"""
        return content_type in self.supported_types